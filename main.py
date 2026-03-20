"""
Analyseur DCE — Backend FastAPI
Déploiement : Railway / Render / local
"""
import os, io, json, zipfile, tempfile
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
import anthropic
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
import fitz  # PyMuPDF

app = FastAPI(title="Analyseur DCE")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
MODEL = "claude-sonnet-4-20250514"

SYSTEM = """Tu es un expert en marchés publics BTP, spécialisé dans la réponse aux appels d'offres.
Réponds UNIQUEMENT en JSON valide, sans markdown, sans backticks.
Ne jamais inventer une donnée absente — écrire "Non trouvé".
Toujours citer la pièce source (ex: CCAP art. 4.3)."""

AUTO_PROMPT = """Analyse complète de ce DCE pour répondre à un appel d'offres marché public BTP.
Réponds UNIQUEMENT avec ce JSON :
{
  "acheteur":"","objet":"","lots":"","corps_etat":"","montant_estime":"",
  "type_marche":"","delai":"","date_limite":"","demarrage":"","lieu":"",
  "pieces_disponibles":[],"pieces_manquantes":[],
  "criteres":[{"nom":"","poids":"","sous_criteres":[]}],
  "obligations":{"regime_prix":"","penalites":"","garanties":"","assurances":[],"clauses_sociales":""},
  "technique":{"materiaux_imposes":[],"performances_exigees":[],"points_risque":[]},
  "dpgf_postes":[{"libelle":"","unite":"","qte":"","risque":"rouge|orange|jaune","commentaire":""}],
  "risques_bloquants":[],"risques_vigilance":[],"points_a_clarifier":[],
  "questions_ao":[{"texte":"","piece":""}],
  "recommandation":"GO|NO-GO|À affiner","recommandation_motif":""
}"""

STEP_PROMPTS = [
    """ÉTAPE 1 - Identité du marché. JSON :
{"acheteur":"","objet":"","lots":"","corps_etat":"","montant_estime":"","type_marche":"","delai":"","date_limite":"","demarrage":"","lieu":"","pieces_disponibles":[],"pieces_manquantes":[]}""",
    """ÉTAPE 2 - Critères de sélection (RC). JSON :
{"criteres":[{"nom":"","poids":"","sous_criteres":[]}],"capacites":{"ca_min":"","references":"","effectifs":"","certifications":""},"atypiques":[]}""",
    """ÉTAPE 3 - Obligations contractuelles (CCAP). JSON :
{"regime_prix":"","revision":"","delai_paiement":"","avance":"","penalites":{"retard":"","plafond":""},"garanties":{"retenue":"","decennale":true},"sous_traitance":"","assurances":[],"clauses_sociales":""}""",
    """ÉTAPE 4 - Exigences techniques (CCTP). JSON :
{"materiaux":[],"performances":[],"methodes":[],"essais":[],"documents":[],"interfaces":[],"risques_tech":[]}""",
    """ÉTAPE 5 - Analyse DPGF/BPU. JSON :
{"postes":[{"libelle":"","unite":"","qte":"","risque":"rouge|orange|jaune","commentaire":""}],"incoherences":[],"hors_marche":[]}""",
    """ÉTAPE 6 - Synthèse risques et recommandation. JSON :
{"risques_bloquants":[],"risques_vigilance":[],"points_a_clarifier":[],"questions_ao":[{"texte":"","piece":""}],"recommandation":"GO|NO-GO|À affiner","recommandation_motif":""}""",
]


# ── Extraction fichiers ──────────────────────────────────────
def extract_pdfs_from_zip(zip_bytes: bytes) -> list[dict]:
    """Extrait récursivement tous les PDF d'un ZIP (sous-dossiers, ZIP imbriqués)."""
    results = []

    def process_zip(zf: zipfile.ZipFile, prefix=""):
        for name in zf.namelist():
            entry = zf.getinfo(name)
            if entry.is_dir():
                continue
            ext = Path(name).suffix.lower()
            data = zf.read(name)
            if ext == ".pdf":
                results.append({"name": prefix + name, "data": data})
            elif ext == ".zip":
                with zipfile.ZipFile(io.BytesIO(data)) as sub:
                    process_zip(sub, prefix + name + "/")
            elif ext in (".doc", ".docx", ".rtf", ".xls", ".xlsx"):
                # Stocker pour info — pas analysable directement
                results.append({"name": prefix + name, "data": None, "skipped": True})

    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
        process_zip(zf)
    return results


def pdf_to_base64_pages(pdf_bytes: bytes) -> list[str]:
    """Convertit un PDF en liste de base64 (1 entrée par document)."""
    import base64
    return [base64.standard_b64encode(pdf_bytes).decode()]


# ── Appel Claude ─────────────────────────────────────────────
def call_claude(pdf_list: list[dict], prompt: str) -> dict:
    content = []
    for pdf in pdf_list[:5]:  # max 5 PDF par appel
        import base64
        content.append({
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf",
                       "data": base64.standard_b64encode(pdf["data"]).decode()},
            "title": pdf["name"],
        })
    content.append({"type": "text", "text": prompt})

    msg = client.messages.create(
        model=MODEL, max_tokens=2000, system=SYSTEM,
        messages=[{"role": "user", "content": content}]
    )
    text = "".join(b.text for b in msg.content if hasattr(b, "text"))
    try:
        return json.loads(text.replace("```json", "").replace("```", "").strip())
    except Exception:
        return {"raw": text, "error": "Parsing JSON échoué"}


# ── Génération Word ──────────────────────────────────────────
def generate_word(data: dict) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_heading(data.get("objet", "Analyse DCE"), 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Acheteur : {data.get('acheteur', '—')}")
    doc.add_paragraph(f"Date limite : {data.get('date_limite', '—')}")
    doc.add_paragraph(f"Montant estimé : {data.get('montant_estime', '—')}")
    doc.add_page_break()

    rec = data.get("recommandation", "")
    motif = data.get("recommandation_motif", "")
    p = doc.add_paragraph()
    r = p.add_run(f"Recommandation : {rec}")
    r.bold = True
    r.font.size = Pt(14)
    if motif:
        doc.add_paragraph(motif)

    sections = [
        ("Critères de sélection", lambda d: [f"{c.get('nom','')} — {c.get('poids','')}" for c in d.get("criteres", [])]),
        ("Risques bloquants 🔴", lambda d: d.get("risques_bloquants", [])),
        ("Points de vigilance 🟠", lambda d: d.get("risques_vigilance", [])),
        ("À clarifier 🟡", lambda d: d.get("points_a_clarifier", [])),
        ("Questions à l'acheteur", lambda d: [q.get("texte", q) if isinstance(q, dict) else q for q in d.get("questions_ao", [])]),
    ]
    for titre, getter in sections:
        items = getter(data)
        if items:
            doc.add_heading(titre, level=1)
            for item in items:
                doc.add_paragraph(str(item), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Génération Excel ─────────────────────────────────────────
def generate_excel(data: dict) -> bytes:
    wb = Workbook()
    BL, WH = "FF1F4E79", "FFFFFFFF"

    def hdr(cell, bg=BL):
        cell.font = Font(bold=True, color=WH)
        cell.fill = PatternFill("solid", fgColor=bg)
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    def flag(v):
        return {"rouge": "FFFF0000", "orange": "FFFF8C00", "jaune": "FFFFFF00"}.get(v.lower() if v else "", "FFD9D9D9")

    # Synthèse
    ws = wb.active
    ws.title = "Synthèse"
    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 50
    rows = [("Champ", "Valeur"), ("Acheteur", data.get("acheteur","")),
            ("Objet", data.get("objet","")), ("Lots", data.get("lots","")),
            ("Montant estimé", data.get("montant_estime","")),
            ("Date limite", data.get("date_limite","")),
            ("Recommandation", data.get("recommandation","")),
            ("Motif", data.get("recommandation_motif",""))]
    for i, (a, b) in enumerate(rows, 1):
        ws.cell(i, 1, a); ws.cell(i, 2, b)
        if i == 1: hdr(ws.cell(i,1)); hdr(ws.cell(i,2))

    # DPGF
    ws2 = wb.create_sheet("DPGF_analyse")
    for j, (w, h) in enumerate([(5,6),(45,20),(10,10),(10,10),(12,12)], 1):
        ws2.column_dimensions[get_column_letter(j)].width = w
    for j, h in enumerate(["#","Libellé","Unité","Qté","Risque"], 1):
        hdr(ws2.cell(1,j)); ws2.cell(1,j,h)
    postes = data.get("dpgf_postes", data.get("postes", []))
    for i, p in enumerate(postes, 2):
        ws2.cell(i,1,i-1); ws2.cell(i,2,p.get("libelle",""))
        ws2.cell(i,3,p.get("unite","")); ws2.cell(i,4,p.get("qte",""))
        niv = p.get("risque","")
        c = ws2.cell(i,5,niv)
        c.fill = PatternFill("solid", fgColor=flag(niv))

    # Obligations
    ws3 = wb.create_sheet("Obligations")
    for j, w in enumerate([20,55,18,12], 1):
        ws3.column_dimensions[get_column_letter(j)].width = w
    for j, h in enumerate(["Thème","Obligation","Source","✓"], 1):
        hdr(ws3.cell(1,j)); ws3.cell(1,j,h)
    obs = data.get("obligations", {})
    rows3 = [("Prix", str(obs.get("regime_prix",""))),
             ("Pénalités", str(obs.get("penalites",""))),
             ("Garanties", str(obs.get("garanties",""))),
             ("Clauses sociales", str(obs.get("clauses_sociales","")))]
    for i, (t, v) in enumerate(rows3, 2):
        ws3.cell(i,1,t); ws3.cell(i,2,v); ws3.cell(i,4,"☐")

    # Questions AO
    ws4 = wb.create_sheet("Questions_AO")
    for j, w in enumerate([5,60,20,18], 1):
        ws4.column_dimensions[get_column_letter(j)].width = w
    for j, h in enumerate(["#","Question","Pièce","Réponse reçue"], 1):
        hdr(ws4.cell(1,j)); ws4.cell(1,j,h)
    for i, q in enumerate(data.get("questions_ao",[]), 2):
        ws4.cell(i,1,i-1)
        ws4.cell(i,2, q.get("texte",q) if isinstance(q,dict) else q)
        ws4.cell(i,3, q.get("piece","") if isinstance(q,dict) else "")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Routes API ───────────────────────────────────────────────
@app.post("/api/analyze")
async def analyze(
    files: list[UploadFile] = File(...),
    mode: str = Form("auto"),
    output: str = Form("word"),
    step: int = Form(0),
    previous_data: str = Form("{}"),
):
    pdfs = []
    skipped = []

    for f in files:
        data = await f.read()
        ext = Path(f.filename).suffix.lower()
        if ext == ".zip":
            extracted = extract_pdfs_from_zip(data)
            for e in extracted:
                if e.get("skipped"):
                    skipped.append(e["name"])
                else:
                    pdfs.append({"name": e["name"], "data": e["data"]})
        elif ext == ".pdf":
            pdfs.append({"name": f.filename, "data": data})
        else:
            skipped.append(f.filename)

    if not pdfs:
        raise HTTPException(400, "Aucun PDF analysable trouvé")

    prompt = AUTO_PROMPT if mode == "auto" else STEP_PROMPTS[min(step, 5)]
    result = call_claude(pdfs, prompt)
    result["_skipped_files"] = skipped
    result["_pdf_count"] = len(pdfs)

    # Générer les fichiers de sortie
    output_files = {}
    merged = {**json.loads(previous_data), **result}

    if output in ("word", "both"):
        output_files["word"] = generate_word(merged).hex()
    if output in ("excel", "both"):
        output_files["excel"] = generate_excel(merged).hex()

    return JSONResponse({"result": result, "files": output_files, "skipped": skipped})


@app.get("/api/download/{fmt}/{hex_data}")
async def download(fmt: str, hex_data: str):
    data = bytes.fromhex(hex_data)
    if fmt == "word":
        return FileResponse(
            path=_write_tmp(data, ".docx"),
            filename="analyse_dce.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif fmt == "excel":
        return FileResponse(
            path=_write_tmp(data, ".xlsx"),
            filename="pilotage_ao.xlsx",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    raise HTTPException(404)

def _write_tmp(data: bytes, ext: str) -> str:
    f = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    f.write(data); f.close()
    return f.name


# ── Servir le front ──────────────────────────────────────────
if os.path.exists("static"):
    app.mount("/", StaticFiles(directory="static", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=int(os.environ.get("PORT", 8000)), reload=True)
